from __future__ import annotations

try:
    from cleaning import clean_text, remove_repeated_headers_footers
except ImportError:
    from src.ingestion.cleaning import clean_text, remove_repeated_headers_footers

import io
import json
import re
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any

import fitz  # PyMuPDF
import numpy as np
from PIL import Image
import easyocr
from docx import Document as DocxDocument


@dataclass
class PageRecord:
    source_id: str
    file_name: str
    file_type: str
    page_number: int
    section: str | None
    timestamp_or_version: str | None
    domain_category: str | None
    raw_text: str
    cleaned_text: str
    image_count: int
    has_numeric_content: bool
    metadata: dict[str, Any]


def detect_numeric_content(text: str) -> bool:
    return bool(re.search(r"\b\d+(?:[.,]\d+)?%?\b", text))


def extract_images_from_page(
    page: fitz.Page,
    output_dir: Path,
    page_number: int,
    file_stem: str,
    source_id: str,
) -> int:
    output_dir.mkdir(parents=True, exist_ok=True)
    page_dict = page.get_text("dict")
    image_blocks = [b for b in page_dict.get("blocks", []) if b.get("type") == 1]

    saved = 0
    safe_stem = re.sub(r"[^a-zA-Z0-9_-]+", "_", file_stem)

    for idx, block in enumerate(image_blocks):
        image_bytes = block.get("image")
        ext = block.get("ext", "png")
        if not image_bytes:
            continue

        try:
            img = Image.open(io.BytesIO(image_bytes))
            img_path = output_dir / (
                f"{source_id}_{safe_stem}_page_{page_number + 1:03d}_img_{idx + 1:02d}.{ext}"
            )
            img.save(img_path)
            saved += 1
        except Exception:
            continue

    return saved


def pil_image_to_numpy(img: Image.Image) -> np.ndarray:
    return np.array(img.convert("RGB"))


def ocr_page_with_easyocr(page: fitz.Page, reader: easyocr.Reader) -> str:
    """
    Render PDF page to image and run EasyOCR.
    """
    pix = page.get_pixmap(dpi=300)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    img_array = pil_image_to_numpy(img)

    result = reader.readtext(img_array, detail=0, paragraph=False)

    extracted_lines = []
    for line in result:
        if isinstance(line, str) and line.strip():
            extracted_lines.append(line.strip())

    return "\n".join(extracted_lines)


def extract_page_text(
    page: fitz.Page,
    ocr_reader: easyocr.Reader | None = None,
    ocr_if_needed: bool = False,
) -> str:
    """
    First try embedded PDF text.
    If missing or too short, fall back to EasyOCR.
    """
    text = page.get_text("text", sort=True)

    if ocr_if_needed and len(text.strip()) < 30 and ocr_reader is not None:
        print(f"Page {page.number + 1}: no embedded text found, trying EasyOCR...")
        try:
            text = ocr_page_with_easyocr(page, ocr_reader)
            print(f"Page {page.number + 1}: EasyOCR extracted {len(text.strip())} characters")
        except Exception as e:
            print(f"Page {page.number + 1}: EasyOCR failed -> {e}")

    return text


def write_records(records: list[PageRecord], output_jsonl: str | Path) -> None:
    output_jsonl = Path(output_jsonl)
    output_jsonl.parent.mkdir(parents=True, exist_ok=True)

    print(f"Writing {len(records)} record(s) to {output_jsonl}")

    with output_jsonl.open("a", encoding="utf-8") as f:
        for rec in records:
            f.write(json.dumps(asdict(rec), ensure_ascii=False) + "\n")

    line_count = sum(1 for _ in output_jsonl.open("r", encoding="utf-8"))
    print(f"File now contains {line_count} line(s)")


def ingest_pdf(
    pdf_path: str | Path,
    output_jsonl: str | Path,
    image_dir: str | Path,
    source_id: str,
    section: str | None = None,
    timestamp_or_version: str | None = None,
    domain_category: str | None = None,
    ocr_if_needed: bool = True,
    ocr_reader: easyocr.Reader | None = None,
) -> list[PageRecord]:
    pdf_path = Path(pdf_path)
    image_dir = Path(image_dir)

    doc = fitz.open(pdf_path)
    print(f"Opened PDF: {pdf_path.name} with {len(doc)} page(s)")

    raw_page_texts: list[str] = []
    page_meta: list[dict[str, Any]] = []
    page_images: list[int] = []

    for page_number, page in enumerate(doc):
        text = extract_page_text(
            page,
            ocr_reader=ocr_reader,
            ocr_if_needed=ocr_if_needed,
        )
        print(f"Page {page_number + 1}: extracted text length = {len(text.strip())}")
        raw_page_texts.append(text)

        img_count = extract_images_from_page(
            page=page,
            output_dir=image_dir,
            page_number=page_number,
            file_stem=pdf_path.stem,
            source_id=source_id,
        )
        page_images.append(img_count)

        page_meta.append(
            {
                "page_width": page.rect.width,
                "page_height": page.rect.height,
                "rotation": page.rotation,
                "file_type": "pdf",
            }
        )

    cleaned_page_texts = remove_repeated_headers_footers(raw_page_texts)

    records: list[PageRecord] = []
    for i, text in enumerate(cleaned_page_texts):
        cleaned = clean_text(text)

        record = PageRecord(
            source_id=source_id,
            file_name=pdf_path.name,
            file_type="pdf",
            page_number=i + 1,
            section=section,
            timestamp_or_version=timestamp_or_version,
            domain_category=domain_category,
            raw_text=raw_page_texts[i],
            cleaned_text=cleaned,
            image_count=page_images[i],
            has_numeric_content=detect_numeric_content(cleaned),
            metadata=page_meta[i],
        )
        records.append(record)

    write_records(records, output_jsonl)
    return records


def extract_docx_text(docx_path: str | Path) -> str:
    """
    Extract paragraphs + tables from DOCX.
    """
    doc = DocxDocument(docx_path)
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def ingest_docx(
    docx_path: str | Path,
    output_jsonl: str | Path,
    source_id: str,
    section: str | None = None,
    timestamp_or_version: str | None = None,
    domain_category: str | None = None,
) -> list[PageRecord]:
    docx_path = Path(docx_path)
    raw_text = extract_docx_text(docx_path)
    cleaned = clean_text(raw_text)

    record = PageRecord(
        source_id=source_id,
        file_name=docx_path.name,
        file_type="docx",
        page_number=1,
        section=section,
        timestamp_or_version=timestamp_or_version,
        domain_category=domain_category,
        raw_text=raw_text,
        cleaned_text=cleaned,
        image_count=0,
        has_numeric_content=detect_numeric_content(cleaned),
        metadata={
            "file_type": "docx",
            "page_width": None,
            "page_height": None,
            "rotation": 0,
        },
    )

    records = [record]
    write_records(records, output_jsonl)
    return records


def read_txt_file(txt_path: str | Path) -> str:
    txt_path = Path(txt_path)

    encodings_to_try = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
    for encoding in encodings_to_try:
        try:
            return txt_path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue

    # last fallback
    return txt_path.read_text(encoding="utf-8", errors="ignore")


def ingest_txt(
    txt_path: str | Path,
    output_jsonl: str | Path,
    source_id: str,
    section: str | None = None,
    timestamp_or_version: str | None = None,
    domain_category: str | None = None,
) -> list[PageRecord]:
    txt_path = Path(txt_path)
    raw_text = read_txt_file(txt_path)
    cleaned = clean_text(raw_text)

    record = PageRecord(
        source_id=source_id,
        file_name=txt_path.name,
        file_type="txt",
        page_number=1,
        section=section,
        timestamp_or_version=timestamp_or_version,
        domain_category=domain_category,
        raw_text=raw_text,
        cleaned_text=cleaned,
        image_count=0,
        has_numeric_content=detect_numeric_content(cleaned),
        metadata={
            "file_type": "txt",
            "page_width": None,
            "page_height": None,
            "rotation": 0,
        },
    )

    records = [record]
    write_records(records, output_jsonl)
    return records


def ingest_file(
    file_path: str | Path,
    output_jsonl: str | Path,
    image_dir: str | Path,
    source_id: str,
    section: str | None = "full_document",
    timestamp_or_version: str | None = "v1",
    domain_category: str | None = "enterprise_document",
    ocr_if_needed: bool = True,
    ocr_reader: easyocr.Reader | None = None,
) -> list[PageRecord]:
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return ingest_pdf(
            pdf_path=file_path,
            output_jsonl=output_jsonl,
            image_dir=image_dir,
            source_id=source_id,
            section=section,
            timestamp_or_version=timestamp_or_version,
            domain_category=domain_category,
            ocr_if_needed=ocr_if_needed,
            ocr_reader=ocr_reader,
        )

    if suffix == ".docx":
        return ingest_docx(
            docx_path=file_path,
            output_jsonl=output_jsonl,
            source_id=source_id,
            section=section,
            timestamp_or_version=timestamp_or_version,
            domain_category=domain_category,
        )

    if suffix == ".txt":
        return ingest_txt(
            txt_path=file_path,
            output_jsonl=output_jsonl,
            source_id=source_id,
            section=section,
            timestamp_or_version=timestamp_or_version,
            domain_category=domain_category,
        )

    raise ValueError(f"Unsupported file type: {file_path.suffix}")


if __name__ == "__main__":
    raw_dir = Path(r"D:\30day git\raw_data_preprocessing_for_LLM\data\raw")
    output_jsonl = Path(r"D:\30day git\raw_data_preprocessing_for_LLM\data\extracted\pages.jsonl")
    image_dir = Path(r"D:\30day git\raw_data_preprocessing_for_LLM\data\extracted\images")

    supported_suffixes = {".pdf", ".docx", ".txt"}

    if output_jsonl.exists():
        output_jsonl.unlink()

    all_files = sorted(
        [p for p in raw_dir.iterdir() if p.is_file() and p.suffix.lower() in supported_suffixes]
    )

    if not all_files:
        print("No supported files found in data/raw")
    else:
        print(f"Found {len(all_files)} supported file(s)")

        needs_ocr = any(p.suffix.lower() == ".pdf" for p in all_files)
        ocr_reader = None

        if needs_ocr:
            print("Loading EasyOCR model once for all PDFs...")
            ocr_reader = easyocr.Reader(["en"], gpu=False)

        all_records: list[PageRecord] = []

        for idx, file_path in enumerate(all_files, start=1):
            source_id = f"doc_{idx:03d}"
            print(f"\nProcessing file: {file_path.name}")

            try:
                records = ingest_file(
                    file_path=file_path,
                    output_jsonl=output_jsonl,
                    image_dir=image_dir,
                    source_id=source_id,
                    section="full_document",
                    timestamp_or_version="v1",
                    domain_category="enterprise_document",
                    ocr_if_needed=True,
                    ocr_reader=ocr_reader,
                )
                all_records.extend(records)
                print(f"Just wrote records for source_id={source_id}, file={file_path.name}")
                print(f"Finished {file_path.name}: {len(records)} record(s)")
            except Exception as e:
                print(f"Failed to process {file_path.name} -> {e}")

        print(f"\nTotal documents processed: {len(all_files)}")
        print(f"Total records written: {len(all_records)}")
        print(f"Output saved to: {output_jsonl}")